# -*- coding: utf-8 -*-
"""TalkTrack-Master v0.1 system TTS validation script.

This script parses a DOCX talktrack, creates a new Shandian admin test IVR when
a valid token is provided, writes normal nodes / jump nodes / end nodes with
system TTS, inserts system-TTS knowledge-base answers, and archives redacted
readback evidence. It reads the backend token only from an environment variable.
"""

from __future__ import annotations

import argparse
import copy
import datetime as dt
import hashlib
import json
import os
import re
import sys
import time
import uuid
from pathlib import Path
from typing import Any
from urllib.parse import unquote

import requests
from docx import Document


BACKENDS = {
    "domestic": {
        "webBase": "https://ai.sd6g.com:1904",
        "apiBase": "https://ai.sd6g.com:1904/api/web",
    },
    "overseas": {
        "webBase": "https://ai.tbot360.com",
        "apiBase": "https://ai.tbot360.com/api/web",
    },
}
TEMPLATE_IVR_ID = 3471
DEFAULT_TTS_VOICE_ID = 1
DEFAULT_SPEECH_RATE = 1
DEFAULT_INDUSTRY_ID = 42
DEFAULT_VAULT_REPORT_DIR = Path(r"D:\ObsidianVault\闪电智能知识库\30-SOP\话术配置")
DEFAULT_ATTACHMENT_DIR = Path(r"D:\ObsidianVault\闪电智能知识库\Attachments\话术配置")
TTS_PATH_RE = re.compile(r"^tts/\d{4}-\d{2}-\d{2}/.+\.wav$")
MODEL_INTENT_PROMPT = """# 角色设定：你是一名客服助手，专注于准确识别用户消息的意图，以支持生成针对性的询问话术回应。

# 任务：根据语义相关性，从[可用意图]中选出与[用户回复]最相关的一个意图。语义相关性衡量[用户回复]与[可用意图]在语义层面的关联，包括关键词匹配、语境契合度和意图表达相似性等因素。相关性越高，表明[用户回复]越可能对应该意图。

# 输入：
客服说的话：「客服：」开头的内容
用户回复：最后一句「用户：」开头的对话内容

# 可用意图：
节点意图：{nodeIntentList}
知识库意图：{knowledgeIntentList}

# 注意：
- “兜底”不是可输出的意图，只是系统在没有命中明确意图后的默认路由。
- 如果[用户回复]无法匹配到明确意图，不得输出“兜底”作为意图；应让系统默认兜底路由处理。
- 如果[用户回复]包含“机器”、“机器人”、“真人”等词，不得识别为语音助手相关意图。
- 当且仅当[用户回复]表示用户不回答时，才可匹配无声意图。
- 只基于[用户回复]分析意图，避免受[客服说的话]影响。

# 输出格式：
只输出结果，禁止输出其他内容，输出示例：{resultFormat}
"""
DEFAULT_MODEL_ID = 55
DEFAULT_MODEL_NAME = "闪电26BMoE-fast"
KEYPAD_PORT_GROUP = "keypadPort"
DEFAULT_MODEL_RESULT_FORMAT = json.dumps([{"intentName": "肯定/默认"}], ensure_ascii=False, separators=(",", ":"))


def now_stamp() -> str:
    return dt.datetime.now().strftime("%Y%m%d_%H%M%S")


def today_slug() -> str:
    return dt.datetime.now().strftime("%Y%m%d")


def sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def ok_code(data: dict[str, Any]) -> bool:
    return str(data.get("code")) == "0"


def extract_access_token(value: str) -> str:
    decoded = unquote((value or "").strip())
    match = re.search(r"[0-9a-fA-F]{32}", decoded)
    if match:
        return match.group(0)
    decoded = re.sub(r"^\s*token\s*=\s*", "", decoded, flags=re.I)
    decoded = re.sub(r"^\s*Bearer\s+", "", decoded, flags=re.I).strip().strip("\"'")
    match = re.search(r"[A-Za-z0-9._-]{20,}", decoded)
    if match:
        return match.group(0)
    raise RuntimeError("No usable access token found. Paste a token, token=Bearer%20..., or -H 'token: Bearer ...'.")


def backend_from_url(value: str) -> str | None:
    lower = (value or "").lower()
    if "ai.sd6g.com:1904" in lower or "sd6g.com:1904" in lower:
        return "domestic"
    if "ai.tbot360.com" in lower or "tbot360.com" in lower:
        return "overseas"
    return None


def probe_backend(region: str, token: str) -> dict[str, Any]:
    meta = BACKENDS[region]
    headers = {"token": f"Bearer {token}", "X-Requested-With": "XMLHttpRequest"}
    try:
        response = requests.get(f"{meta['apiBase']}/account/findInfo", headers=headers, timeout=20)
        response.encoding = "utf-8"
        data = response.json()
    except Exception as exc:
        return {"region": region, "ok": False, "error": type(exc).__name__}
    return {"region": region, "ok": ok_code(data), "code": data.get("code"), "data": data.get("data") or {}}


def resolve_backend(token_text: str, backend_region: str = "auto", backend_url: str = "") -> dict[str, Any]:
    token = extract_access_token(token_text)
    hinted = backend_from_url(backend_url)
    if backend_region != "auto" and hinted and backend_region != hinted:
        raise RuntimeError(f"backend region conflict: --backend-region={backend_region}, --backend-url points to {hinted}")
    if backend_region != "auto":
        candidates = [backend_region]
    elif hinted:
        candidates = [hinted]
    else:
        candidates = list(BACKENDS)
    probes = [probe_backend(region, token) for region in candidates]
    ok = [item for item in probes if item.get("ok")]
    if not ok:
        raise RuntimeError(f"token validation failed for candidate backends: {probes}")
    if len(ok) > 1:
        raise RuntimeError("token validates against multiple backends; pass --backend-region domestic or overseas")
    region = ok[0]["region"]
    meta = BACKENDS[region]
    return {"region": region, "token": token, "webBase": meta["webBase"], "apiBase": meta["apiBase"], "accountInfo": ok[0]}


def safe_json_loads(value: Any, default: Any) -> Any:
    if value is None:
        return default
    if isinstance(value, (dict, list)):
        return value
    if isinstance(value, str) and value.strip():
        return json.loads(value)
    return default


def load_model_intent_config(value: Any) -> dict[str, Any]:
    try:
        data = safe_json_loads(value, {})
    except Exception:
        return {}
    return data if isinstance(data, dict) else {}


def model_config_id(value: Any) -> Any:
    cfg = load_model_intent_config(value)
    if not cfg:
        return None
    model_cfg = cfg.get("modelConfig")
    if isinstance(model_cfg, dict):
        return model_cfg.get("id") if model_cfg.get("id") is not None else model_cfg.get("modelId")
    return cfg.get("id") if cfg.get("id") is not None else cfg.get("modelId")


def is_default_model_id(value: Any) -> bool:
    try:
        return int(value) == DEFAULT_MODEL_ID
    except (TypeError, ValueError):
        return False


def default_keypad_port_group() -> dict[str, Any]:
    return {
        "position": {"name": "absolute"},
        "attrs": {
            "circle": {
                "r": 5,
                "magnet": True,
                "stroke": "#1677ff",
                "strokeWidth": 2,
                "fill": "#fff",
            }
        },
    }


def graph_cell_port_summary(cell: dict[str, Any] | None) -> dict[str, Any]:
    if not cell:
        return {
            "exists": False,
            "hasKeypadPortGroup": False,
            "usesGenericInOutPortGroups": False,
            "portItemGroups": [],
            "allItemsUseKeypadPort": False,
            "portItemIds": [],
            "hasPosition": False,
            "hasSize": False,
            "shape": None,
        }
    ports = cell.get("ports") if isinstance(cell.get("ports"), dict) else {}
    groups = ports.get("groups") if isinstance(ports.get("groups"), dict) else {}
    items = ports.get("items") if isinstance(ports.get("items"), list) else []
    item_groups = sorted({str(item.get("group") or "") for item in items if isinstance(item, dict)})
    item_ids = [str(item.get("id")) for item in items if isinstance(item, dict) and item.get("id") is not None]
    return {
        "exists": True,
        "hasKeypadPortGroup": KEYPAD_PORT_GROUP in groups,
        "usesGenericInOutPortGroups": "in" in groups or "out" in groups or "in" in item_groups or "out" in item_groups,
        "portItemGroups": item_groups,
        "allItemsUseKeypadPort": bool(items) and all(isinstance(item, dict) and item.get("group") == KEYPAD_PORT_GROUP for item in items),
        "portItemIds": item_ids,
        "hasPosition": isinstance(cell.get("position"), dict),
        "hasSize": isinstance(cell.get("size"), dict),
        "shape": cell.get("shape"),
    }


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def normalize_spaces(text: str) -> str:
    text = (text or "").replace("\u3000", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def clean_spoken_text(text: str) -> str:
    text = normalize_spaces(text)
    text = text.replace("】", "")
    text = re.sub(r"^[□☆]\s*", "", text)
    text = re.sub(r"^(肯定|默认|拒绝|否定|任何)[、/，]?.*?[:：]\s*", "", text)
    text = re.sub(r"^[-—]+", "", text)
    text = text.replace("——", " ")
    text = text.replace("（挂机）", "").replace("(挂机)", "")
    text = re.sub(r"，?并给客户打标签.*$", "", text)
    text = re.sub(r"。?\s*——挂机.*$", "", text)
    text = normalize_spaces(text)
    return text


def text_after_marker(text: str, marker: str) -> str:
    if marker not in text:
        return clean_spoken_text(text)
    return clean_spoken_text(text.split(marker, 1)[1])


def split_keywords(raw: str) -> list[str]:
    raw = normalize_spaces(raw)
    raw = re.sub(r"^【[^】]+】", "", raw).strip()
    pieces = re.split(r"[；;、,，?？/]+", raw)
    seen: set[str] = set()
    result: list[str] = []
    for piece in pieces:
        item = normalize_spaces(piece)
        item = item.strip("；;、,，?？ ")
        if not item or item in seen:
            continue
        seen.add(item)
        result.append(item)
    return result


def title_from_question(question: str, fallback: str) -> str:
    match = re.match(r"^【([^】]+)】", normalize_spaces(question))
    if match:
        return match.group(1)
    keywords = split_keywords(question)
    return (keywords[0] if keywords else fallback)[:24]


def first_answer_segment(answer: str) -> str:
    answer = normalize_spaces(answer)
    answer = re.split(r"\s*/\s*[□☆]", answer, maxsplit=1)[0]
    answer = re.split(r"\s*[□☆]\s*(肯定|默认|拒绝|否定|任何)", answer, maxsplit=1)[0]
    return clean_spoken_text(answer)


def next_nonempty_after(paragraphs: list[str], heading: str) -> str:
    for index, paragraph in enumerate(paragraphs):
        if paragraph.startswith(heading):
            for item in paragraphs[index + 1 :]:
                if item and not item.startswith(("□", "☆")) and not item.startswith("流程"):
                    return clean_spoken_text(item)
    return ""


def collect_flow2(paragraphs: list[str]) -> str:
    result: list[str] = []
    in_flow2 = False
    for paragraph in paragraphs:
        if paragraph.startswith("流程2"):
            in_flow2 = True
            continue
        if paragraph.startswith("流程3"):
            break
        if not in_flow2:
            continue
        if paragraph.startswith(("□", "☆")):
            continue
        if paragraph:
            result.append(clean_spoken_text(paragraph))
    return normalize_spaces(" ".join(result))


def extract_branch_text(paragraphs: list[str], contains: str, marker: str) -> str:
    for paragraph in paragraphs:
        if contains in paragraph and marker in paragraph:
            return text_after_marker(paragraph, marker)
    return ""


def parse_docx(docx_path: Path, max_kb: int) -> dict[str, Any]:
    doc = Document(str(docx_path))
    paragraphs = [normalize_spaces(p.text) for p in doc.paragraphs if normalize_spaces(p.text)]
    flow1_opening = next_nonempty_after(paragraphs, "流程1")
    flow1_recovery = extract_branch_text(paragraphs, "挽回", "挽回")
    flow2_intro = collect_flow2(paragraphs)
    flow3_invite = next_nonempty_after(paragraphs, "流程3")
    hangup = extract_branch_text(paragraphs, "挂机", "挂机")
    if not hangup:
        hangup = "抱歉，打扰了，这边先不继续占用您的时间，再见。"

    all_candidates: list[dict[str, Any]] = []
    group_names = {
        1: "活动业务问题",
        2: "项目业务问题",
        3: "一般问题",
    }
    for table_index, table in enumerate(doc.tables, start=1):
        group = group_names.get(table_index, f"表格{table_index}")
        for row_index, row in enumerate(table.rows, start=1):
            cells = [normalize_spaces(cell.text.replace("\n", " ")) for cell in row.cells]
            if row_index == 1 or len(cells) < 2:
                continue
            question, answer = cells[0], cells[1]
            if not question or question == "问题" or not answer:
                continue
            answer_text = first_answer_segment(answer)
            if not answer_text:
                continue
            keywords = split_keywords(question)
            all_candidates.append(
                {
                    "sourceGroup": group,
                    "sourceTable": table_index,
                    "sourceRow": row_index,
                    "title": title_from_question(question, f"{group}{row_index}"),
                    "keyword": ",".join(keywords[:12]),
                    "keywordCount": len(keywords),
                    "rawQuestion": question,
                    "answerText": answer_text,
                    "answerHash": sha256_text(answer_text)[:16],
                }
            )

    selected = select_kb_candidates(all_candidates, max_kb)
    return {
        "sourcePath": str(docx_path),
        "sourceFileName": docx_path.name,
        "paragraphCount": len(doc.paragraphs),
        "tableCount": len(doc.tables),
        "normalNodes": [
            {"key": "flow1_opening", "name": "流程1-开场白", "spokenText": flow1_opening},
            {"key": "flow1_recovery", "name": "挽回-二次介绍", "spokenText": flow1_recovery},
            {"key": "flow2_intro", "name": "流程2-活动介绍", "spokenText": flow2_intro},
        ],
        "jumpNodes": [
            {"key": "jump_to_flow2", "name": "跳转-进入活动介绍", "targetKey": "flow2_intro"},
        ],
        "endNodes": [
            {"key": "flow3_invite_end", "name": "邀约成功-结束", "spokenText": flow3_invite},
            {"key": "reject_hangup_end", "name": "拒绝挂机-结束", "spokenText": hangup},
        ],
        "branches": [
            {"from": "flow1_opening", "intent": "肯定/默认", "to": "jump_to_flow2"},
            {"from": "flow1_opening", "intent": "拒绝/否定", "to": "flow1_recovery"},
            {"from": "flow1_recovery", "intent": "肯定/默认", "to": "jump_to_flow2"},
            {"from": "flow1_recovery", "intent": "拒绝/否定", "to": "reject_hangup_end"},
            {"from": "flow2_intro", "intent": "肯定/默认", "to": "flow3_invite_end"},
            {"from": "flow2_intro", "intent": "拒绝/否定", "to": "reject_hangup_end"},
        ],
        "knowledgeBaseCandidateCount": len(all_candidates),
        "knowledgeBaseCandidatesSelectedCount": len(selected),
        "knowledgeBaseCandidatesSelected": selected,
        "knowledgeBaseCandidateGroups": group_counts(all_candidates),
        "warnings": parse_warnings(flow1_opening, flow1_recovery, flow2_intro, flow3_invite, hangup, selected),
    }


def select_kb_candidates(candidates: list[dict[str, Any]], max_kb: int) -> list[dict[str, Any]]:
    if max_kb <= 0:
        return []
    by_group: dict[str, list[dict[str, Any]]] = {}
    for item in candidates:
        by_group.setdefault(item["sourceGroup"], []).append(item)
    group_order = ["活动业务问题", "项目业务问题", "一般问题"]
    selected: list[dict[str, Any]] = []
    cursor = 0
    while len(selected) < max_kb:
        added = False
        for group in group_order:
            items = by_group.get(group) or []
            if cursor < len(items):
                selected.append(items[cursor])
                added = True
                if len(selected) >= max_kb:
                    break
        if not added:
            break
        cursor += 1
    return selected


def group_counts(candidates: list[dict[str, Any]]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for item in candidates:
        counts[item["sourceGroup"]] = counts.get(item["sourceGroup"], 0) + 1
    return counts


def parse_warnings(*values: Any) -> list[str]:
    warnings: list[str] = []
    for index, value in enumerate(values, start=1):
        if isinstance(value, str) and not value:
            warnings.append(f"解析字段 {index} 为空")
    return warnings


class Client:
    def __init__(self, token: str, backend: dict[str, Any]):
        self.backend = backend
        self.base_url = backend["apiBase"]
        self.web_base = backend["webBase"]
        self.headers = {
            "token": f"Bearer {token}",
            "X-Requested-With": "XMLHttpRequest",
        }

    def url(self, path: str) -> str:
        return f"{self.base_url}{path}"

    def script_graph_url(self, ivr_id: int) -> str:
        return f"{self.web_base}/script-graph?ivrId={ivr_id}"

    def get(self, path: str, timeout: int = 60) -> dict[str, Any]:
        response = requests.get(self.url(path), headers=self.headers, timeout=timeout)
        return self._json(response, path)

    def post(self, path: str, body: Any | None = None, timeout: int = 90) -> dict[str, Any]:
        response = requests.post(self.url(path), headers=self.headers, json=body, timeout=timeout)
        return self._json(response, path)

    def post_url(self, url: str, body: Any | None = None, timeout: int = 90) -> dict[str, Any]:
        response = requests.post(url, headers=self.headers, json=body, timeout=timeout)
        return self._json(response, url)

    @staticmethod
    def _json(response: requests.Response, path: str) -> dict[str, Any]:
        response.encoding = "utf-8"
        try:
            data = response.json()
        except Exception as exc:
            raise RuntimeError(f"{path} returned non-JSON HTTP {response.status_code}: {response.text[:300]}") from exc
        if response.status_code >= 400:
            raise RuntimeError(f"{path} HTTP {response.status_code}: {json.dumps(data, ensure_ascii=False)[:500]}")
        return data

    @staticmethod
    def assert_ok(data: dict[str, Any], action: str) -> None:
        if not ok_code(data):
            raise RuntimeError(f"{action} failed: {json.dumps(data, ensure_ascii=False)[:1000]}")


def make_tts_item(text: str, record_file_path: str) -> dict[str, Any]:
    item_id = uuid.uuid4().hex
    return {
        "id": item_id,
        "type": 2,
        "name": text,
        "recordFilePath": record_file_path,
        "paramId": -1,
        "paramIdList": [],
        "contentList": [],
    }


def make_model_intent_config(result_label: str = "肯定/默认", model_id: int = DEFAULT_MODEL_ID) -> dict[str, Any]:
    return {
        "modelTimeoutMilliSecond": 2000,
        "modelConfig": {
            "id": model_id,
            "prompt": MODEL_INTENT_PROMPT,
            "resultFormat": json.dumps([{"intentName": result_label}], ensure_ascii=False, separators=(",", ":")),
            "temperature": 0.0,
            "topK": 0,
            "topP": 0.0001,
            "presencePenalty": 0.0,
            "maxTokens": 4096,
            "thinkBudget": 1,
            "seed": 0,
            "enableThinking": 0,
            "recognitionRound": 1,
        },
    }


def normalize_model_intent_config(node: dict[str, Any], result_label: str = "肯定/默认") -> None:
    template = make_model_intent_config(result_label)
    cfg = load_model_intent_config(node.get("modelIntentRecognitionConfig"))
    if not cfg:
        node["modelIntentRecognitionConfig"] = template
        return

    cfg.setdefault("modelTimeoutMilliSecond", template["modelTimeoutMilliSecond"])
    if not isinstance(cfg.get("modelConfig"), dict):
        cfg["modelConfig"] = {}
    model_cfg = cfg["modelConfig"]
    for key, value in template["modelConfig"].items():
        if model_cfg.get(key) in (None, ""):
            model_cfg[key] = value
    # Node-level 2.0 must always use 闪电26BMoE-fast. Existing prompts and
    # result formats are preserved, but stale model selections are corrected.
    model_cfg["id"] = DEFAULT_MODEL_ID
    if "modelId" in model_cfg:
        model_cfg["modelId"] = DEFAULT_MODEL_ID
    node["modelIntentRecognitionConfig"] = cfg


def enable_model_intent_2(node: dict[str, Any], result_label: str = "肯定/默认") -> None:
    node["modelIntentRecognitionEnabled"] = 1
    normalize_model_intent_config(node, result_label)


def make_ivr_model_intent_payload(
    ivr_id: int,
    prompt: str = MODEL_INTENT_PROMPT,
    result_format: str = DEFAULT_MODEL_RESULT_FORMAT,
) -> dict[str, Any]:
    return {
        "id": ivr_id,
        "modelIntentRecognitionEnabled": 1,
        "modelIntentRecognitionTimeoutMilliSecond": 2000,
        "modelPrompt": prompt,
        "modelId": DEFAULT_MODEL_ID,
        "modelResultFormat": result_format,
        "modelTemperature": 0,
        "modelTopP": 0,
        "modelTopK": 0,
        "modelPresencePenalty": 0,
        "modelMaxTokens": 4096,
        "modelThinkBudget": 0,
        "modelSeed": 0,
        "modelEnableThinking": 0,
        "modelRecognitionRound": 0,
    }


def update_ivr_model_intent_recognition_config(client: Client, ivr_id: int) -> dict[str, Any]:
    payload = make_ivr_model_intent_payload(ivr_id)
    model_intent_config_url = client.url("/ivr/updateModelIntentRecognitionConfig")
    updated = client.post_url(model_intent_config_url, payload, timeout=120)
    Client.assert_ok(updated, "update IVR model intent recognition 2.0")
    return {
        "updateModelIntentRecognitionConfig": updated,
        "modelIntentRecognitionConfigUrl": model_intent_config_url,
        "modelIntentRecognitionPayloadSummary": {
            "id": ivr_id,
            "modelIntentRecognitionEnabled": 1,
            "modelIntentRecognitionTimeoutMilliSecond": 2000,
            "modelId": DEFAULT_MODEL_ID,
            "modelMaxTokens": 4096,
            "modelRecognitionRound": 0,
            "modelPromptSha256": sha256_text(payload["modelPrompt"]),
            "modelPromptLength": len(payload["modelPrompt"]),
            "modelResultFormat": payload["modelResultFormat"],
        },
    }


def has_intent(node: dict[str, Any], intent_id: str) -> bool:
    for item in node.get("intentList") or []:
        if not isinstance(item, dict):
            continue
        if str(intent_id) in {str(key) for key in item.keys()}:
            return True
        if str(item.get("value")) == str(intent_id):
            return True
    return False


def intent_target_map_from_backend(node: dict[str, Any]) -> dict[str, str]:
    target_map: dict[str, str] = {}
    for item in node.get("intentList") or []:
        if not isinstance(item, dict) or "value" in item:
            continue
        for intent_id, target_id in item.items():
            if target_id:
                target_map[str(intent_id)] = str(target_id)
    return target_map


def frontend_intent_rows_from_backend(
    backend_node: dict[str, Any],
    label_by_intent: dict[str, str] | None = None,
    existing_sources: list[dict[str, Any] | None] | None = None,
) -> list[dict[str, str]]:
    labels: dict[str, str] = {"-2": "知识库", "-1": "兜底"}
    labels.update(label_by_intent or {})
    digit_sequences: dict[str, str] = {}

    for source in existing_sources or []:
        if not isinstance(source, dict):
            continue
        for row in source.get("intentList") or []:
            if isinstance(row, dict) and row.get("value") is not None:
                value = str(row.get("value"))
                if row.get("label"):
                    labels[value] = str(row.get("label"))
                digit_sequences[value] = str(row.get("digitSequence") or "")
        for row in source.get("ports") or []:
            if isinstance(row, dict) and row.get("value") is not None:
                value = str(row.get("value"))
                if row.get("label"):
                    labels[value] = str(row.get("label"))
                digit_sequences[value] = str(row.get("digitSequence") or "")

    rows: list[dict[str, str]] = []
    for item in backend_node.get("intentList") or []:
        if not isinstance(item, dict) or "value" in item:
            continue
        for intent_id in item.keys():
            value = str(intent_id)
            rows.append(
                {
                    "value": value,
                    "label": labels.get(value, value),
                    "digitSequence": digit_sequences.get(value, ""),
                }
            )
    return rows


def frontend_node_copy_from_backend(backend_node: dict[str, Any], intent_rows: list[dict[str, str]]) -> dict[str, Any]:
    node = copy.deepcopy(backend_node)
    node["intentList"] = copy.deepcopy(intent_rows)
    return node


def ensure_knowledge_base_intent(node: dict[str, Any]) -> None:
    if has_intent(node, "-2"):
        return
    intents = node.setdefault("intentList", [])
    insert_at = len(intents)
    for index, item in enumerate(intents):
        if isinstance(item, dict) and "-1" in {str(key) for key in item.keys()}:
            insert_at = index
            break
    intents.insert(insert_at, {"-2": ""})


def enable_knowledge_base_matching(node: dict[str, Any], knowledge_base_ids: list[int] | None = None) -> None:
    ensure_knowledge_base_intent(node)
    node["matchKnowledgeBaseEnabled"] = 1
    node["knowledgeBaseMatchType"] = 2
    node["knowledgeBaseMatchList"] = list(knowledge_base_ids or [])
    node["notMatchedKnowledgeBaseList"] = []


def strip_template_refs(node: dict[str, Any]) -> None:
    node["interruptedIntentList"] = []
    node["dtmfIntentMappingList"] = []
    node["dtmfIntentList"] = []
    node["llmNodeIntentMappingList"] = []
    node["llmNodeIntentList"] = []
    node["customIntentTagList"] = []
    node["customIntentTagListJson"] = "[]"
    node["intentMatchFilterEnabled"] = 0
    node["intentMatchFilterMilliSecond"] = 0
    node["intentMatchWaitEnabled"] = 0
    node["intentMatchWaitMilliSecond"] = 0
    node["intentUnknownWaitEnabled"] = 0
    node["intentUnknownWaitMilliSecond"] = 0
    node["intentWaitEnabled"] = 0
    node["intentWaitRuleList"] = []
    node["modelIntentRecognitionEnabled"] = 0
    node.pop("modelIntentRecognitionConfig", None)
    node["setUserIntentLevelEnabled"] = 0
    node["setUserIntentLevelId"] = None
    node["userIntentLevelName"] = None
    node["knowledgeBaseMatchList"] = []
    node["notMatchedKnowledgeBaseList"] = []


def synthesize_text(client: Client, ivr_id: int, text: str, voice_id: int, speech_rate: float) -> dict[str, Any]:
    body = {"ivrId": ivr_id, "ttsVoiceId": voice_id, "speechRate": speech_rate, "text": text}
    create = client.post("/ivr/createNodeTextTtsRecord", body, timeout=120)
    query_results: list[dict[str, Any]] = []
    final_data: dict[str, Any] | None = None
    for _ in range(18):
        query = client.post("/ivr/queryTtsRecord", body, timeout=120)
        query_results.append({"code": query.get("code"), "message": query.get("message"), "data": query.get("data")})
        data = query.get("data") or {}
        if ok_code(query) and data.get("recordFilePath"):
            final_data = data
            break
        time.sleep(1)
    if not final_data:
        final_data = (query_results[-1].get("data") if query_results else {}) or {}
    path = final_data.get("recordFilePath")
    return {
        "text": text,
        "textHash": sha256_text(text)[:16],
        "createCode": create.get("code"),
        "createMessage": create.get("message"),
        "queryCode": query_results[-1].get("code") if query_results else None,
        "queryMessage": query_results[-1].get("message") if query_results else None,
        "synthesisStatus": final_data.get("synthesisStatus"),
        "recordFilePath": path,
        "pathLooksLikeSystemTts": isinstance(path, str) and bool(TTS_PATH_RE.match(path)),
    }


def set_system_tts_node(node: dict[str, Any], name: str, text: str, tts_path: str, *, node_type: int) -> None:
    strip_template_refs(node)
    node["name"] = name
    node["text"] = text
    node["type"] = node_type
    node["recordType"] = 2
    node["recordFilePath"] = ""
    node["textParamIncluded"] = 0
    node["recordPlaybackList"] = []
    node["recordPlaybackListJson"] = "[]"
    tts_list = [make_tts_item(text, tts_path)]
    node["ttsPlaybackList"] = tts_list
    node["ttsPlaybackListJson"] = json.dumps(tts_list, ensure_ascii=False, separators=(",", ":"))
    node["allowInterruptEnabled"] = 1 if node_type == 1 else 0
    node["allowInterruptSecond"] = 2 if node_type == 1 else 0
    if node_type == 1:
        enable_knowledge_base_matching(node)
        enable_model_intent_2(node)
    else:
        node["matchKnowledgeBaseEnabled"] = 0
    if node_type == 2:
        node["nextType"] = 2
        node["nextSceneId"] = None
        node["nextNodeId"] = None
        node["intentList"] = []


def set_jump_node(node: dict[str, Any], name: str, target_scene_id: str, target_node_id: str) -> None:
    strip_template_refs(node)
    node["name"] = name
    node["text"] = ""
    node["type"] = 2
    node["recordType"] = 1
    node["recordFilePath"] = ""
    node["recordPlaybackList"] = []
    node["recordPlaybackListJson"] = "[]"
    node["ttsPlaybackList"] = []
    node["ttsPlaybackListJson"] = "[]"
    node["intentList"] = []
    node["allowInterruptEnabled"] = 0
    node["allowInterruptSecond"] = 0
    node["matchKnowledgeBaseEnabled"] = 0
    node["nextType"] = 1
    node["nextSceneId"] = target_scene_id
    node["nextNodeId"] = target_node_id


def find_node(nodes: list[dict[str, Any]], node_id: str) -> dict[str, Any]:
    for node in nodes:
        if node.get("id") == node_id:
            return node
    raise KeyError(node_id)


def rebuild_ports(existing: dict[str, Any], rows: list[dict[str, str]], target_map: dict[str, str], width: float, height: float) -> dict[str, Any]:
    groups = copy.deepcopy(existing.get("groups") or {})
    groups.pop("in", None)
    groups.pop("out", None)
    if KEYPAD_PORT_GROUP not in groups:
        groups[KEYPAD_PORT_GROUP] = default_keypad_port_group()
    items = []
    count = max(len(rows), 1)
    for index, row in enumerate(rows):
        port_id = str(row["value"])
        items.append(
            {
                "id": port_id,
                "group": KEYPAD_PORT_GROUP,
                "args": {"x": width, "y": round((index + 1) * height / (count + 1), 2)},
                "attrs": {"text": {"text": row.get("label", port_id)}},
                "data": {"targetNodeId": target_map.get(port_id, "")},
            }
        )
    return {"groups": groups, "items": items}


def sync_frontend(front_scene: dict[str, Any], backend_node: dict[str, Any], label_by_intent: dict[str, str]) -> None:
    node_id = backend_node.get("id")
    existing_front_node: dict[str, Any] | None = None
    existing_cell_data: dict[str, Any] | None = None
    for index, node in enumerate(front_scene.get("nodeList") or []):
        if node.get("id") == node_id:
            existing_front_node = node
            break

    for cell in (front_scene.get("graph") or {}).get("cells") or []:
        if cell.get("id") == node_id:
            existing_cell_data = cell.get("data") or {}
            break

    port_rows = frontend_intent_rows_from_backend(
        backend_node,
        label_by_intent,
        [existing_front_node, existing_cell_data, (existing_cell_data or {}).get("customData")],
    )
    target_map = intent_target_map_from_backend(backend_node)
    frontend_node = frontend_node_copy_from_backend(backend_node, port_rows)

    for index, node in enumerate(front_scene.get("nodeList") or []):
        if node.get("id") == node_id:
            front_scene["nodeList"][index] = copy.deepcopy(frontend_node)
            break

    for cell in (front_scene.get("graph") or {}).get("cells") or []:
        if cell.get("id") != node_id:
            continue
        data = cell.setdefault("data", {})
        data["label"] = backend_node.get("name")
        data["title"] = backend_node.get("name")
        data["description"] = backend_node.get("text") or ""
        data["nodeType"] = backend_node.get("type")
        if backend_node.get("type") == 2 and backend_node.get("nextType") == 1:
            data["actionName"] = "跳转"
        elif backend_node.get("type") == 2 and backend_node.get("nextType") == 2:
            data["actionName"] = "挂机"
        data["customData"] = copy.deepcopy(frontend_node)
        data["ports"] = port_rows
        size = cell.get("size") or {}
        cell["ports"] = rebuild_ports(cell.get("ports") or {}, port_rows, target_map, size.get("width", 276), size.get("height", 176))


def build_edges(front_scene: dict[str, Any], backend_nodes: list[dict[str, Any]]) -> None:
    graph = front_scene.setdefault("graph", {})
    cells = graph.get("cells") or []
    edge_template = next(
        (copy.deepcopy(cell) for cell in cells if cell.get("source") or cell.get("target") or cell.get("shape") == "edge"),
        {"shape": "edge", "attrs": {"line": {"stroke": "#000", "strokeWidth": "2"}}, "zIndex": 1},
    )
    node_ids = {node.get("id") for node in backend_nodes}
    non_edges = [cell for cell in cells if not (cell.get("source") or cell.get("target") or cell.get("shape") == "edge")]
    edges = []
    z_index = 100
    for node in backend_nodes:
        for item in node.get("intentList") or []:
            for intent_id, target_id in item.items():
                if target_id and target_id in node_ids:
                    edge = copy.deepcopy(edge_template)
                    edge["id"] = str(uuid.uuid4())
                    edge["zIndex"] = z_index
                    edge["source"] = {"cell": node["id"], "port": str(intent_id)}
                    edge["target"] = {"cell": target_id}
                    edges.append(edge)
                    z_index += 1
    graph["cells"] = non_edges + edges
    front_scene["graph"] = graph


def resolve_ivr_id(client: Client, created: dict[str, Any], name: str) -> int:
    data = created.get("data")
    if isinstance(data, int):
        return data
    if isinstance(data, str) and data.isdigit():
        return int(data)
    if isinstance(data, dict):
        for key in ("id", "ivrId"):
            if data.get(key):
                return int(data[key])
    page = client.post("/ivr/findPage", {"query": {"searchName": name}, "page": {"current": 1, "size": 20}})
    Client.assert_ok(page, "find created IVR")
    rows = (page.get("data") or {}).get("list") or (page.get("data") or {}).get("records") or []
    for row in rows:
        if row.get("name") == name:
            return int(row["id"])
    raise RuntimeError("Could not resolve created IVR id.")


def create_test_ivr(client: Client, name: str, voice_id: int, speech_rate: float) -> tuple[int, dict[str, Any]]:
    created = client.post(
        "/ivr/insert",
        {
            "voiceType": 2,
            "ttsVoiceId": voice_id,
            "speechRate": speech_rate,
            "name": name,
            "industryId": DEFAULT_INDUSTRY_ID,
        },
    )
    Client.assert_ok(created, "create test IVR")
    ivr_id = resolve_ivr_id(client, created, name)
    update = client.post(
        "/ivr/update",
        {
            "id": ivr_id,
            "name": name,
            "voiceType": 2,
            "ttsVoiceId": voice_id,
            "speechRate": speech_rate,
            "industryId": DEFAULT_INDUSTRY_ID,
        },
    )
    return ivr_id, {"insert": created, "update": update}


def get_intents(client: Client, ivr_id: int) -> dict[str, str]:
    data = client.get(f"/ivrIntent/findList/{ivr_id}")
    Client.assert_ok(data, "read intents")
    result: dict[str, str] = {}
    for item in data.get("data") or []:
        name = str(item.get("name") or "")
        item_id = str(item.get("id") or "")
        if not item_id:
            continue
        if "肯定" in name or "默认" in name:
            result.setdefault("positive", item_id)
        if "否定" in name or "拒绝" in name:
            result.setdefault("negative", item_id)
        if "不回答" in name:
            result.setdefault("no_answer", item_id)
    missing = [key for key in ("positive", "negative") if key not in result]
    if missing:
        raise RuntimeError(f"Missing default intents in new IVR: {missing}")
    return result


def configure_scene(client: Client, ivr_id: int, parsed: dict[str, Any], voice_id: int, speech_rate: float) -> tuple[dict[str, Any], list[dict[str, Any]], dict[str, Any]]:
    template = client.get(f"/ivr/findSceneList/{TEMPLATE_IVR_ID}", timeout=120)
    Client.assert_ok(template, "read template scene")
    scene_list = copy.deepcopy(safe_json_loads(template["data"]["sceneList"], []))
    front_list = copy.deepcopy(safe_json_loads(template["data"]["sceneListFrontend"], []))
    if not scene_list or not front_list:
        raise RuntimeError("Template scene is empty.")

    scene = scene_list[0]
    front_scene = front_list[0]
    scene["name"] = "活动邀约系统TTS验证主场景"
    front_scene["name"] = scene["name"]
    scene_id = scene["id"]
    nodes = scene.get("nodeList") or []
    normal_nodes = [node for node in nodes if int(node.get("type") or 0) == 1]
    action_nodes = [node for node in nodes if int(node.get("type") or 0) == 2]
    if len(normal_nodes) < 3 or len(action_nodes) < 3:
        raise RuntimeError("Template scene does not have enough normal/action nodes.")

    node_by_key = {
        "flow1_opening": normal_nodes[0],
        "flow1_recovery": normal_nodes[1],
        "flow2_intro": normal_nodes[2],
        "jump_to_flow2": action_nodes[0],
        "flow3_invite_end": action_nodes[1],
        "reject_hangup_end": action_nodes[2],
    }

    tts_results: list[dict[str, Any]] = []
    for node_spec in parsed["normalNodes"]:
        node = node_by_key[node_spec["key"]]
        tts = synthesize_text(client, ivr_id, node_spec["spokenText"], voice_id, speech_rate)
        tts_results.append({"scope": "normal_node", "nodeKey": node_spec["key"], "nodeId": node.get("id"), "nodeName": node_spec["name"], **tts})
        set_system_tts_node(node, node_spec["name"], node_spec["spokenText"], tts.get("recordFilePath") or "", node_type=1)

    for node_spec in parsed["endNodes"]:
        node = node_by_key[node_spec["key"]]
        tts = synthesize_text(client, ivr_id, node_spec["spokenText"], voice_id, speech_rate)
        tts_results.append({"scope": "end_node", "nodeKey": node_spec["key"], "nodeId": node.get("id"), "nodeName": node_spec["name"], **tts})
        set_system_tts_node(node, node_spec["name"], node_spec["spokenText"], tts.get("recordFilePath") or "", node_type=2)

    set_jump_node(
        node_by_key["jump_to_flow2"],
        "跳转-进入活动介绍",
        scene_id,
        node_by_key["flow2_intro"]["id"],
    )

    intents = get_intents(client, ivr_id)
    positive = intents["positive"]
    negative = intents["negative"]
    label_by_intent = {positive: "肯定/默认", negative: "拒绝/否定", "-2": "知识库", "-1": "兜底"}
    node_by_key["flow1_opening"]["intentList"] = [{positive: node_by_key["jump_to_flow2"]["id"]}, {negative: node_by_key["flow1_recovery"]["id"]}, {"-2": ""}, {"-1": node_by_key["flow1_recovery"]["id"]}]
    node_by_key["flow1_recovery"]["intentList"] = [{positive: node_by_key["jump_to_flow2"]["id"]}, {negative: node_by_key["reject_hangup_end"]["id"]}, {"-2": ""}, {"-1": node_by_key["reject_hangup_end"]["id"]}]
    node_by_key["flow2_intro"]["intentList"] = [{positive: node_by_key["flow3_invite_end"]["id"]}, {negative: node_by_key["reject_hangup_end"]["id"]}, {"-2": ""}, {"-1": node_by_key["reject_hangup_end"]["id"]}]

    for node in nodes:
        strip_template_refs(node)
        if node.get("id") not in {n.get("id") for n in node_by_key.values()}:
            node["intentList"] = []
            node["matchKnowledgeBaseEnabled"] = 0
        if int(node.get("type") or 0) == 1 and node.get("intentList"):
            enable_knowledge_base_matching(node)
            enable_model_intent_2(node, label_by_intent.get(positive, "肯定/默认"))
        sync_frontend(front_scene, node, label_by_intent)
    build_edges(front_scene, nodes)

    payload = {
        "ivrId": ivr_id,
        "sceneList": json.dumps([scene], ensure_ascii=False, separators=(",", ":")),
        "sceneListFrontend": json.dumps([front_scene], ensure_ascii=False, separators=(",", ":")),
    }
    updated = client.post("/ivr/updateSceneList", payload, timeout=120)
    Client.assert_ok(updated, "update scene list")
    readback = client.get(f"/ivr/findSceneList/{ivr_id}", timeout=120)
    Client.assert_ok(readback, "read back scene")
    return readback, tts_results, {"updateSceneList": updated, "intentIds": intents, "nodeIdByKey": {key: node.get("id") for key, node in node_by_key.items()}, "sceneId": scene_id}


def create_kb(client: Client, ivr_id: int, item: dict[str, Any], priority: int, voice_id: int, speech_rate: float) -> dict[str, Any]:
    tts = synthesize_text(client, ivr_id, item["answerText"], voice_id, speech_rate)
    tts_item = make_tts_item(item["answerText"], tts.get("recordFilePath") or "")
    payload = {
        "ivrId": ivr_id,
        "type": 1 if item["sourceGroup"] != "项目业务问题" else 2,
        "title": item["title"],
        "keyword": item["keyword"],
        "priority": priority,
        "popupEnabled": 0,
        "setUserIntentLevelEnabled": 0,
        "sendSmsEnabled": 0,
        "textParamIncluded": 0,
        "recordType": 2,
        "answerLoopMethod": 1,
        "matchBeforeIntentEnabled": 0,
        "customIntentTagListJson": "[]",
        "takeoverAfterPlaybackEnabled": 0,
        "addBlacklistEnabled": 0,
        "matchSpeechInterimEnabled": 0,
        "matchSpeechInterimSpeakRoundLimit": 0,
        "answerList": [
            {
                "text": item["answerText"],
                "nextType": 0,
                "nextSceneId": None,
                "nextNodeId": None,
                "nextAfterUserSpeechEnabled": 0,
                "nextAfterInterruptEnabled": 0,
                "nextTypeAfterInterrupt": 0,
                "nextSceneIdAfterInterrupt": None,
                "nextNodeIdAfterInterrupt": None,
                "recordType": 2,
                "recordFilePath": "",
                "textParamIncluded": 0,
                "recordPlaybackListJson": "[]",
                "ttsPlaybackListJson": json.dumps([tts_item], ensure_ascii=False, separators=(",", ":")),
                "answerOrder": 1,
            }
        ],
    }
    inserted = client.post("/ivrKnowledgeBase/insert", payload, timeout=120)
    Client.assert_ok(inserted, f"insert KB {item['title']}")
    return {
        "sourceGroup": item["sourceGroup"],
        "sourceRow": item["sourceRow"],
        "title": item["title"],
        "keyword": item["keyword"],
        "answerText": item["answerText"],
        "priority": priority,
        "insertCode": inserted.get("code"),
        "insertMessage": inserted.get("message"),
        "tts": tts,
    }


def create_knowledge_bases(client: Client, ivr_id: int, parsed: dict[str, Any], voice_id: int, speech_rate: float) -> tuple[list[dict[str, Any]], dict[str, Any], list[dict[str, Any]]]:
    before = client.post("/ivrKnowledgeBase/findPage", {"query": {"ivrId": ivr_id, "searchName": ""}, "page": {"current": 1, "size": 200}}, timeout=120)
    Client.assert_ok(before, "read existing KB")
    rows = (before.get("data") or {}).get("records") or (before.get("data") or {}).get("list") or []
    start_priority = max([int(row.get("priority") or 0) for row in rows] or [-1]) + 1
    inserts = [
        create_kb(client, ivr_id, item, start_priority + index, voice_id, speech_rate)
        for index, item in enumerate(parsed["knowledgeBaseCandidatesSelected"])
    ]
    page = client.post("/ivrKnowledgeBase/findPage", {"query": {"ivrId": ivr_id, "searchName": ""}, "page": {"current": 1, "size": 200}}, timeout=120)
    Client.assert_ok(page, "read KB page")
    rows = (page.get("data") or {}).get("records") or (page.get("data") or {}).get("list") or []
    by_title = {row.get("title"): row for row in rows}
    details = []
    for item in parsed["knowledgeBaseCandidatesSelected"]:
        row = by_title.get(item["title"])
        if not row:
            details.append({"title": item["title"], "found": False})
            continue
        detail = client.post(f"/ivrKnowledgeBase/findById/{row.get('id')}", None, timeout=120)
        Client.assert_ok(detail, f"read KB detail {item['title']}")
        details.append({"title": item["title"], "found": True, "listRow": row, "detail": detail.get("data")})
    return inserts, page, details


def read_knowledge_base_ids(client: Client, ivr_id: int) -> list[int]:
    page = client.post("/ivrKnowledgeBase/findPage", {"query": {"ivrId": ivr_id, "searchName": ""}, "page": {"current": 1, "size": 1000}}, timeout=120)
    Client.assert_ok(page, "read KB ids")
    rows = (page.get("data") or {}).get("records") or (page.get("data") or {}).get("list") or []
    return [int(row["id"]) for row in rows if row.get("id") is not None]


def sync_matching_to_frontend(front_list: list[dict[str, Any]], backend_node: dict[str, Any]) -> None:
    node_id = backend_node.get("id")
    for scene in front_list:
        existing_front_node: dict[str, Any] | None = None
        existing_cell_data: dict[str, Any] | None = None
        for index, node in enumerate(scene.get("nodeList") or []):
            if node.get("id") == node_id:
                existing_front_node = node
                break
        for cell in (scene.get("graph") or {}).get("cells") or []:
            if cell.get("id") == node_id:
                existing_cell_data = cell.get("data") or {}
                break
        rows = frontend_intent_rows_from_backend(
            backend_node,
            None,
            [existing_front_node, existing_cell_data, (existing_cell_data or {}).get("customData")],
        )
        frontend_node = frontend_node_copy_from_backend(backend_node, rows)
        for index, node in enumerate(scene.get("nodeList") or []):
            if node.get("id") == node_id:
                scene["nodeList"][index] = copy.deepcopy(frontend_node)
                break
        target_map = intent_target_map_from_backend(backend_node)
        for cell in (scene.get("graph") or {}).get("cells") or []:
            if cell.get("id") != node_id:
                continue
            data = cell.setdefault("data", {})
            data["customData"] = copy.deepcopy(frontend_node)
            data["ports"] = copy.deepcopy(rows)
            size = cell.get("size") or {}
            cell["ports"] = rebuild_ports(cell.get("ports") or {}, rows, target_map, size.get("width", 276), size.get("height", 176))


def apply_knowledge_base_matching(client: Client, ivr_id: int) -> tuple[dict[str, Any], dict[str, Any]]:
    knowledge_base_ids = read_knowledge_base_ids(client, ivr_id)
    current = client.get(f"/ivr/findSceneList/{ivr_id}", timeout=120)
    Client.assert_ok(current, "read scene before KB matching")
    scene_list = safe_json_loads(current["data"].get("sceneList"), [])
    front_list = safe_json_loads(current["data"].get("sceneListFrontend"), [])
    changed_nodes = []
    for scene in scene_list:
        for node in scene.get("nodeList") or []:
            if int(node.get("type") or 0) != 1 or not node.get("intentList"):
                continue
            before = {
                "matchKnowledgeBaseEnabled": node.get("matchKnowledgeBaseEnabled"),
                "knowledgeBaseMatchType": node.get("knowledgeBaseMatchType"),
                "knowledgeBaseMatchListCount": len(node.get("knowledgeBaseMatchList") or []),
                "hasKnowledgeBaseIntent": has_intent(node, "-2"),
            }
            enable_knowledge_base_matching(node, knowledge_base_ids)
            after = {
                "matchKnowledgeBaseEnabled": node.get("matchKnowledgeBaseEnabled"),
                "knowledgeBaseMatchType": node.get("knowledgeBaseMatchType"),
                "knowledgeBaseMatchListCount": len(node.get("knowledgeBaseMatchList") or []),
                "hasKnowledgeBaseIntent": has_intent(node, "-2"),
            }
            if before != after:
                changed_nodes.append({"id": node.get("id"), "name": node.get("name"), "before": before, "after": after})
            sync_matching_to_frontend(front_list, node)
    payload = {
        "ivrId": ivr_id,
        "sceneList": json.dumps(scene_list, ensure_ascii=False, separators=(",", ":")),
        "sceneListFrontend": json.dumps(front_list, ensure_ascii=False, separators=(",", ":")),
    }
    updated = client.post("/ivr/updateSceneList", payload, timeout=120)
    Client.assert_ok(updated, "update KB matching")
    readback = client.get(f"/ivr/findSceneList/{ivr_id}", timeout=120)
    Client.assert_ok(readback, "read back KB matching")
    return readback, {"updateSceneListForKnowledgeBaseMatching": updated, "knowledgeBaseIds": knowledge_base_ids, "changedNodes": changed_nodes}


def parse_tts_json(value: Any) -> list[dict[str, Any]]:
    parsed = safe_json_loads(value, [])
    return parsed if isinstance(parsed, list) else []


def validate_run(ivr_id: int, parsed: dict[str, Any], scene_readback: dict[str, Any], node_tts: list[dict[str, Any]], kb_tts: list[dict[str, Any]], kb_details: list[dict[str, Any]], scene_meta: dict[str, Any]) -> dict[str, Any]:
    root_data = scene_readback.get("data") or {}
    scene_list = safe_json_loads(root_data.get("sceneList"), [])
    front_list = safe_json_loads(root_data.get("sceneListFrontend"), [])
    nodes = [node for scene in scene_list for node in scene.get("nodeList") or []]
    front_nodes = [node for scene in front_list for node in scene.get("nodeList") or []]
    nodes_by_id = {node.get("id"): node for node in nodes}
    front_nodes_by_id = {node.get("id"): node for node in front_nodes}
    graph_cells = []
    for scene in front_list:
        graph_cells.extend((scene.get("graph") or {}).get("cells") or [])
    graph_by_id = {cell.get("id"): cell for cell in graph_cells if cell.get("id")}
    root_model_intent_summary = {
        "updateApiCode": (scene_meta.get("updateModelIntentRecognitionConfig") or {}).get("code"),
        "modelIntentRecognitionEnabled": root_data.get("modelIntentRecognitionEnabled"),
        "modelIntentRecognitionTimeoutMilliSecond": root_data.get("modelIntentRecognitionTimeoutMilliSecond"),
        "modelId": root_data.get("modelId"),
        "modelPromptPresent": bool(root_data.get("modelPrompt")),
        "modelPromptLength": len(root_data.get("modelPrompt") or ""),
        "modelResultFormatPresent": bool(root_data.get("modelResultFormat")),
        "modelTemperature": root_data.get("modelTemperature"),
        "modelTopP": root_data.get("modelTopP"),
        "modelTopK": root_data.get("modelTopK"),
        "modelPresencePenalty": root_data.get("modelPresencePenalty"),
        "modelMaxTokens": root_data.get("modelMaxTokens"),
        "modelThinkBudget": root_data.get("modelThinkBudget"),
        "modelSeed": root_data.get("modelSeed"),
        "modelEnableThinking": root_data.get("modelEnableThinking"),
        "modelRecognitionRound": root_data.get("modelRecognitionRound"),
        "payloadSummary": scene_meta.get("modelIntentRecognitionPayloadSummary") or {},
    }

    node_summary = []
    for key, node_id in scene_meta["nodeIdByKey"].items():
        node = nodes_by_id.get(node_id) or {}
        front_node = front_nodes_by_id.get(node_id) or {}
        graph_cell = graph_by_id.get(node_id) or {}
        graph_custom = (((graph_cell or {}).get("data") or {}).get("customData") or {})
        if not isinstance(graph_custom, dict):
            graph_custom = {}
        port_summary = graph_cell_port_summary(graph_cell)
        tts_list = node.get("ttsPlaybackList") or []
        node_summary.append(
            {
                "nodeKey": key,
                "nodeId": node_id,
                "nodeName": node.get("name"),
                "type": node.get("type"),
                "recordType": node.get("recordType"),
                "ttsPlaybackListCount": len(tts_list),
                "ttsRecordFilePath": (tts_list[0].get("recordFilePath") if tts_list else None),
                "nextType": node.get("nextType"),
                "nextSceneId": node.get("nextSceneId"),
                "nextNodeId": node.get("nextNodeId"),
                "intentList": node.get("intentList") or [],
                "hasKnowledgeBaseIntent": has_intent(node, "-2"),
                "matchKnowledgeBaseEnabled": node.get("matchKnowledgeBaseEnabled"),
                "knowledgeBaseMatchType": node.get("knowledgeBaseMatchType"),
                "knowledgeBaseMatchListCount": len(node.get("knowledgeBaseMatchList") or []),
                "modelIntentRecognitionEnabled": node.get("modelIntentRecognitionEnabled"),
                "modelIntentRecognitionConfigPresent": bool(node.get("modelIntentRecognitionConfig")),
                "modelIntentRecognitionModelId": model_config_id(node.get("modelIntentRecognitionConfig")),
                "frontendNodeExists": node_id in front_nodes_by_id,
                "graphCellExists": node_id in graph_by_id,
                "frontendHasKnowledgeBaseIntent": has_intent(front_node, "-2"),
                "frontendMatchKnowledgeBaseEnabled": front_node.get("matchKnowledgeBaseEnabled"),
                "frontendKnowledgeBaseMatchType": front_node.get("knowledgeBaseMatchType"),
                "frontendKnowledgeBaseMatchListCount": len(front_node.get("knowledgeBaseMatchList") or []),
                "graphHasKnowledgeBaseIntent": has_intent(graph_custom, "-2"),
                "graphMatchKnowledgeBaseEnabled": graph_custom.get("matchKnowledgeBaseEnabled"),
                "graphKnowledgeBaseMatchType": graph_custom.get("knowledgeBaseMatchType"),
                "graphKnowledgeBaseMatchListCount": len(graph_custom.get("knowledgeBaseMatchList") or []),
                "frontendModelIntentRecognitionEnabled": front_node.get("modelIntentRecognitionEnabled"),
                "frontendModelIntentRecognitionConfigPresent": bool(front_node.get("modelIntentRecognitionConfig")),
                "frontendModelIntentRecognitionModelId": model_config_id(front_node.get("modelIntentRecognitionConfig")),
                "graphModelIntentRecognitionEnabled": graph_custom.get("modelIntentRecognitionEnabled"),
                "graphModelIntentRecognitionConfigPresent": bool(graph_custom.get("modelIntentRecognitionConfig")),
                "graphModelIntentRecognitionModelId": model_config_id(graph_custom.get("modelIntentRecognitionConfig")),
                "graphCustomDataNextNodeId": graph_custom.get("nextNodeId"),
                "graphHasKeypadPortGroup": port_summary.get("hasKeypadPortGroup"),
                "graphUsesGenericInOutPortGroups": port_summary.get("usesGenericInOutPortGroups"),
                "graphPortItemGroups": port_summary.get("portItemGroups"),
                "graphAllItemsUseKeypadPort": port_summary.get("allItemsUseKeypadPort"),
                "graphPortItemIds": port_summary.get("portItemIds"),
                "graphHasPosition": port_summary.get("hasPosition"),
                "graphHasSize": port_summary.get("hasSize"),
                "graphShape": port_summary.get("shape"),
            }
        )

    kb_summary = []
    for item in kb_details:
        row = item.get("listRow") or {}
        detail = item.get("detail") or {}
        answers = detail.get("answerList") or row.get("answerList") or []
        answer_summary = []
        for answer in answers:
            tts_list = parse_tts_json(answer.get("ttsPlaybackListJson"))
            answer_summary.append(
                {
                    "answerId": answer.get("id"),
                    "text": answer.get("text"),
                    "recordType": answer.get("recordType"),
                    "ttsPlaybackListCount": len(tts_list),
                    "ttsRecordFilePath": tts_list[0].get("recordFilePath") if tts_list else None,
                    "pathLooksLikeSystemTts": bool(tts_list and TTS_PATH_RE.match(str(tts_list[0].get("recordFilePath") or ""))),
                }
            )
        kb_summary.append(
            {
                "knowledgeBaseId": row.get("id") or detail.get("id"),
                "title": item.get("title"),
                "keyword": row.get("keyword") or detail.get("keyword"),
                "found": item.get("found"),
                "priority": row.get("priority") or detail.get("priority"),
                "answers": answer_summary,
            }
        )

    jump_checks = []
    for key in ("jump_to_flow2",):
        jump = nodes_by_id.get(scene_meta["nodeIdByKey"].get(key))
        if not jump:
            jump_checks.append({"nodeKey": key, "ok": False, "reason": "jump node not found"})
            continue
        target_id = jump.get("nextNodeId")
        graph_custom = ((graph_by_id.get(jump.get("id")) or {}).get("data") or {}).get("customData") or {}
        jump_checks.append(
            {
                "nodeKey": key,
                "nodeName": jump.get("name"),
                "nextType": jump.get("nextType"),
                "nextSceneId": jump.get("nextSceneId"),
                "nextNodeId": target_id,
                "targetExists": target_id in nodes_by_id,
                "graphCustomDataMatches": isinstance(graph_custom, dict) and graph_custom.get("nextNodeId") == target_id,
            }
        )

    edge_checks = []
    node_ids = set(nodes_by_id.keys())
    edge_cells = [cell for cell in graph_cells if cell.get("source") or cell.get("target") or cell.get("shape") == "edge"]
    for edge in edge_cells:
        source = edge.get("source") or {}
        target = edge.get("target") or {}
        source_cell = graph_by_id.get(source.get("cell")) or {}
        source_ports = set(graph_cell_port_summary(source_cell).get("portItemIds") or [])
        source_port = str(source.get("port") or "")
        edge_checks.append(
            {
                "sourceCell": source.get("cell"),
                "sourcePort": source.get("port"),
                "targetCell": target.get("cell"),
                "sourceExists": source.get("cell") in node_ids,
                "targetExists": target.get("cell") in node_ids,
                "sourcePortExists": bool(source_port and source_port in source_ports),
            }
        )

    failures: list[str] = []
    if str(root_model_intent_summary.get("updateApiCode")) != "0":
        failures.append("IVR-level model intent recognition 2.0 update API did not return code=0")
    if int(root_model_intent_summary.get("modelIntentRecognitionEnabled") or 0) != 1:
        failures.append("IVR-level Advanced Settings model intent recognition 2.0 is not enabled")
    if int(root_model_intent_summary.get("modelId") or 0) != DEFAULT_MODEL_ID:
        failures.append(f"IVR-level modelId != {DEFAULT_MODEL_ID}")
    if not root_model_intent_summary.get("modelPromptPresent"):
        failures.append("IVR-level modelPrompt is empty or not echoed by readback")
    if root_model_intent_summary.get("modelMaxTokens") is not None and int(root_model_intent_summary.get("modelMaxTokens") or 0) < 4096:
        failures.append("IVR-level modelMaxTokens < 4096")
    if root_model_intent_summary.get("modelRecognitionRound") is not None and int(root_model_intent_summary.get("modelRecognitionRound") or 0) != 0:
        failures.append("IVR-level modelRecognitionRound != 0")
    normal_keys = {item["key"] for item in parsed["normalNodes"]}
    end_keys = {item["key"] for item in parsed["endNodes"]}
    for item in node_summary:
        key = item["nodeKey"]
        backend_uses_model_2 = int(item.get("modelIntentRecognitionEnabled") or 0) == 1 or bool(item.get("modelIntentRecognitionConfigPresent"))
        frontend_uses_model_2 = int(item.get("frontendModelIntentRecognitionEnabled") or 0) == 1 or bool(item.get("frontendModelIntentRecognitionConfigPresent"))
        graph_uses_model_2 = int(item.get("graphModelIntentRecognitionEnabled") or 0) == 1 or bool(item.get("graphModelIntentRecognitionConfigPresent"))
        if backend_uses_model_2 and not is_default_model_id(item.get("modelIntentRecognitionModelId")):
            failures.append(f"node {key} backend 2.0 model is not {DEFAULT_MODEL_NAME}({DEFAULT_MODEL_ID})")
        if frontend_uses_model_2 and not is_default_model_id(item.get("frontendModelIntentRecognitionModelId")):
            failures.append(f"node {key} frontend 2.0 model is not {DEFAULT_MODEL_NAME}({DEFAULT_MODEL_ID})")
        if graph_uses_model_2 and not is_default_model_id(item.get("graphModelIntentRecognitionModelId")):
            failures.append(f"node {key} graph 2.0 model is not {DEFAULT_MODEL_NAME}({DEFAULT_MODEL_ID})")
        routed_graph_node = bool(item.get("intentList")) or int(item.get("nextType") or 0) == 1
        if routed_graph_node and item.get("graphCellExists"):
            if not item.get("graphHasKeypadPortGroup"):
                failures.append(f"node {key} graph ports missing {KEYPAD_PORT_GROUP}")
            if item.get("graphUsesGenericInOutPortGroups"):
                failures.append(f"node {key} graph ports use generic in/out groups instead of {KEYPAD_PORT_GROUP}")
            if item.get("graphPortItemIds") and not item.get("graphAllItemsUseKeypadPort"):
                failures.append(f"node {key} graph port items are not all {KEYPAD_PORT_GROUP}")
            if not item.get("graphHasPosition") or not item.get("graphHasSize"):
                failures.append(f"node {key} graph cell missing position/size render fields")
        if key in normal_keys:
            if int(item.get("type") or 0) != 1:
                failures.append(f"普通节点 {key} type != 1")
            if int(item.get("recordType") or 0) != 2:
                failures.append(f"普通节点 {key} recordType != 2")
            if int(item.get("ttsPlaybackListCount") or 0) <= 0:
                failures.append(f"普通节点 {key} ttsPlaybackList 为空")
            if not TTS_PATH_RE.match(str(item.get("ttsRecordFilePath") or "")):
                failures.append(f"普通节点 {key} TTS 路径不是系统 TTS")
            if int(item.get("matchKnowledgeBaseEnabled") or 0) != 1 or int(item.get("knowledgeBaseMatchType") or 0) != 2 or int(item.get("knowledgeBaseMatchListCount") or 0) <= 0 or not item.get("hasKnowledgeBaseIntent"):
                failures.append(f"normal node {key} NLP knowledge-base matching missing in backend sceneList")
            if int(item.get("frontendMatchKnowledgeBaseEnabled") or 0) != 1 or int(item.get("frontendKnowledgeBaseMatchType") or 0) != 2 or int(item.get("frontendKnowledgeBaseMatchListCount") or 0) <= 0 or not item.get("frontendHasKnowledgeBaseIntent"):
                failures.append(f"normal node {key} NLP knowledge-base matching missing in frontend nodeList")
            if int(item.get("graphMatchKnowledgeBaseEnabled") or 0) != 1 or int(item.get("graphKnowledgeBaseMatchType") or 0) != 2 or int(item.get("graphKnowledgeBaseMatchListCount") or 0) <= 0 or not item.get("graphHasKnowledgeBaseIntent"):
                failures.append(f"normal node {key} NLP knowledge-base matching missing in graph customData")
            if int(item.get("modelIntentRecognitionEnabled") or 0) != 1 or not item.get("modelIntentRecognitionConfigPresent"):
                failures.append(f"normal node {key} model intent recognition 2.0 missing in backend sceneList")
            if int(item.get("frontendModelIntentRecognitionEnabled") or 0) != 1 or not item.get("frontendModelIntentRecognitionConfigPresent"):
                failures.append(f"normal node {key} model intent recognition 2.0 missing in frontend nodeList")
            if int(item.get("graphModelIntentRecognitionEnabled") or 0) != 1 or not item.get("graphModelIntentRecognitionConfigPresent"):
                failures.append(f"normal node {key} model intent recognition 2.0 missing in graph customData")
        if key in end_keys:
            if int(item.get("type") or 0) != 2:
                failures.append(f"结束节点 {key} type != 2")
            if int(item.get("recordType") or 0) != 2:
                failures.append(f"结束节点 {key} recordType != 2")
            if int(item.get("ttsPlaybackListCount") or 0) <= 0:
                failures.append(f"结束节点 {key} ttsPlaybackList 为空")
    for item in node_tts + kb_tts:
        if not item.get("pathLooksLikeSystemTts"):
            failures.append(f"TTS 路径异常: {item.get('scope')} {item.get('nodeName') or item.get('title')}")
    for check in jump_checks:
        if int(check.get("nextType") or 0) != 1:
            failures.append(f"跳转节点 {check.get('nodeKey')} nextType != 1")
        if not check.get("targetExists"):
            failures.append(f"跳转节点 {check.get('nodeKey')} target 不存在")
        if not check.get("graphCustomDataMatches"):
            failures.append(f"跳转节点 {check.get('nodeKey')} graph customData 不一致")
    for edge in edge_checks:
        if not edge.get("sourceExists") or not edge.get("targetExists"):
            failures.append(f"graph edge source/target 不存在: {edge}")
        if edge.get("sourceExists") and not edge.get("sourcePortExists"):
            failures.append(f"graph edge source port 不存在或端口结构异常: {edge}")
    for kb in kb_summary:
        if not kb.get("found"):
            failures.append(f"知识库未读回: {kb.get('title')}")
        for answer in kb.get("answers") or []:
            if int(answer.get("recordType") or 0) != 2:
                failures.append(f"知识库 {kb.get('title')} answer recordType != 2")
            if int(answer.get("ttsPlaybackListCount") or 0) <= 0:
                failures.append(f"知识库 {kb.get('title')} ttsPlaybackListJson 为空")
            if not answer.get("pathLooksLikeSystemTts"):
                failures.append(f"知识库 {kb.get('title')} TTS 路径异常")

    return {
        "passed": not failures,
        "failureItems": failures,
        "rootModelIntentSummary": root_model_intent_summary,
        "nodeSummary": node_summary,
        "jumpChecks": jump_checks,
        "edgeChecks": edge_checks,
        "knowledgeBaseSummary": kb_summary,
    }


def markdown_table(rows: list[list[Any]]) -> str:
    if not rows:
        return ""
    out = ["| " + " | ".join(str(v) for v in rows[0]) + " |"]
    out.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    for row in rows[1:]:
        out.append("| " + " | ".join(str(v).replace("\n", " ") for v in row) + " |")
    return "\n".join(out)


def write_report(report: dict[str, Any], report_path: Path, json_path: Path) -> None:
    validation = report.get("validation") or {}
    parsed = report.get("parsed") or {}
    node_rows = [["节点Key", "节点名", "type", "recordType", "TTS数", "TTS路径", "KB后端", "KB前端", "KB画布", "2.0后端", "2.0前端", "2.0画布", "2.0模型后端", "2.0模型前端", "2.0模型画布", "画布端口", "nextType", "nextNodeId"]]
    for item in validation.get("nodeSummary") or []:
        node_rows.append([
            item.get("nodeKey"),
            item.get("nodeName"),
            item.get("type"),
            item.get("recordType"),
            item.get("ttsPlaybackListCount"),
            item.get("ttsRecordFilePath") or "",
            bool(int(item.get("matchKnowledgeBaseEnabled") or 0) == 1 and int(item.get("knowledgeBaseMatchType") or 0) == 2 and int(item.get("knowledgeBaseMatchListCount") or 0) > 0 and item.get("hasKnowledgeBaseIntent")),
            bool(int(item.get("frontendMatchKnowledgeBaseEnabled") or 0) == 1 and int(item.get("frontendKnowledgeBaseMatchType") or 0) == 2 and int(item.get("frontendKnowledgeBaseMatchListCount") or 0) > 0 and item.get("frontendHasKnowledgeBaseIntent")),
            bool(int(item.get("graphMatchKnowledgeBaseEnabled") or 0) == 1 and int(item.get("graphKnowledgeBaseMatchType") or 0) == 2 and int(item.get("graphKnowledgeBaseMatchListCount") or 0) > 0 and item.get("graphHasKnowledgeBaseIntent")),
            bool(int(item.get("modelIntentRecognitionEnabled") or 0) == 1 and item.get("modelIntentRecognitionConfigPresent")),
            bool(int(item.get("frontendModelIntentRecognitionEnabled") or 0) == 1 and item.get("frontendModelIntentRecognitionConfigPresent")),
            bool(int(item.get("graphModelIntentRecognitionEnabled") or 0) == 1 and item.get("graphModelIntentRecognitionConfigPresent")),
            item.get("modelIntentRecognitionModelId") or "",
            item.get("frontendModelIntentRecognitionModelId") or "",
            item.get("graphModelIntentRecognitionModelId") or "",
            bool(item.get("graphHasKeypadPortGroup") and not item.get("graphUsesGenericInOutPortGroups") and (not item.get("graphPortItemIds") or item.get("graphAllItemsUseKeypadPort"))),
            item.get("nextType") or "",
            item.get("nextNodeId") or "",
        ])
    kb_rows = [["KB ID", "标题", "关键词", "answerId", "recordType", "TTS数", "TTS路径"]]
    for kb in validation.get("knowledgeBaseSummary") or []:
        for answer in kb.get("answers") or []:
            kb_rows.append([kb.get("knowledgeBaseId"), kb.get("title"), kb.get("keyword"), answer.get("answerId"), answer.get("recordType"), answer.get("ttsPlaybackListCount"), answer.get("ttsRecordFilePath") or ""])
    jump_rows = [["节点Key", "节点名", "nextType", "nextSceneId", "nextNodeId", "targetExists", "graphCustomDataMatches"]]
    for item in validation.get("jumpChecks") or []:
        jump_rows.append([item.get("nodeKey"), item.get("nodeName"), item.get("nextType"), item.get("nextSceneId"), item.get("nextNodeId"), item.get("targetExists"), item.get("graphCustomDataMatches")])
    root_model = validation.get("rootModelIntentSummary") or {}
    root_model_rows = [
        ["字段", "读回值"],
        ["updateApiCode", root_model.get("updateApiCode")],
        ["modelIntentRecognitionEnabled", root_model.get("modelIntentRecognitionEnabled")],
        ["modelIntentRecognitionTimeoutMilliSecond", root_model.get("modelIntentRecognitionTimeoutMilliSecond")],
        ["modelId", root_model.get("modelId")],
        ["modelPromptPresent", root_model.get("modelPromptPresent")],
        ["modelPromptLength", root_model.get("modelPromptLength")],
        ["modelResultFormatPresent", root_model.get("modelResultFormatPresent")],
        ["modelMaxTokens", root_model.get("modelMaxTokens")],
        ["modelRecognitionRound", root_model.get("modelRecognitionRound")],
    ]

    failures = validation.get("failureItems") or []
    failure_text = "无" if not failures else "\n".join(f"- {item}" for item in failures)
    status = "通过" if validation.get("passed") else ("未完成" if not report.get("ivrId") else "部分通过")
    md = f"""---
tags:
  - 话术配置
  - TalkTrack-Master
  - 系统TTS
  - 回归测试
created: {dt.datetime.now().strftime('%Y-%m-%d')}
source: [[普通节点配置Skill_v0.1范围与验收_20260512]]
---

# TalkTrack-Master v0.1 系统 TTS 回归报告 {today_slug()}

## 结论

验证结论：{status}

- 测试 IVR：`{report.get('ivrId') or '未创建'}` / `{report.get('ivrName') or '未创建'}`
- 页面地址：{report.get('scriptGraphUrl') or '未创建'}
- 源 DOCX：`{parsed.get('sourcePath')}`
- 使用系统音色：`{report.get('voice', {}).get('ttsVoiceId')}` / `{report.get('voice', {}).get('voiceName')}`
- 语速：`{report.get('voice', {}).get('speechRate')}`
- 脱敏 JSON：[[{json_path.name}]]
- token：仅从当前环境变量读取，未写入脚本、报告、JSON 或截图。

## DOCX 解析摘要

- 段落数：`{parsed.get('paragraphCount')}`
- 表格数：`{parsed.get('tableCount')}`
- 普通节点候选：`{len(parsed.get('normalNodes') or [])}`
- 跳转节点候选：`{len(parsed.get('jumpNodes') or [])}`
- 结束节点候选：`{len(parsed.get('endNodes') or [])}`
- 知识库答案候选总数：`{parsed.get('knowledgeBaseCandidateCount')}`
- 本次回归写入知识库答案样本数：`{parsed.get('knowledgeBaseCandidatesSelectedCount')}`
- 候选分布：`{json.dumps(parsed.get('knowledgeBaseCandidateGroups') or {}, ensure_ascii=False)}`

## 普通节点和结束节点读回

{markdown_table(node_rows)}

## 跳转节点读回

{markdown_table(jump_rows)}

## 话术级大模型意图分析 2.0 读回

{markdown_table(root_model_rows)}

## 知识库答案读回

{markdown_table(kb_rows)}

## 接口结果

| 接口 | 结果 |
| --- | --- |
| `GET /account/findInfo` | `{report.get('apiCodes', {}).get('accountFindInfo')}` |
| `GET /ivr/findAllTtsVoiceBaseInfo` | `{report.get('apiCodes', {}).get('voiceList')}` |
| `POST /ivr/insert` | `{report.get('apiCodes', {}).get('ivrInsert')}` |
| `POST /ivr/update` | `{report.get('apiCodes', {}).get('ivrUpdate')}` |
| `GET /ivr/findSceneList/3471` | `{report.get('apiCodes', {}).get('templateRead')}` |
| `POST /ivr/createNodeTextTtsRecord` | 节点 `{report.get('apiCodes', {}).get('nodeTtsCreateOkCount')}` 次成功；知识库 `{report.get('apiCodes', {}).get('kbTtsCreateOkCount')}` 次成功 |
| `POST /ivr/queryTtsRecord` | 节点 `{report.get('apiCodes', {}).get('nodeTtsQueryOkCount')}` 次成功；知识库 `{report.get('apiCodes', {}).get('kbTtsQueryOkCount')}` 次成功 |
| `POST /ivr/updateSceneList` | `{report.get('apiCodes', {}).get('updateSceneList')}` |
| `POST /ivr/updateSceneList` 知识库匹配补齐 | `{report.get('apiCodes', {}).get('updateSceneListForKnowledgeBaseMatching')}` |
| `POST /ivr/updateModelIntentRecognitionConfig` | `{report.get('apiCodes', {}).get('updateModelIntentRecognitionConfig')}` |
| `POST /ivrKnowledgeBase/insert` | `{report.get('apiCodes', {}).get('kbInsertOkCount')}` 条成功 |
| `POST /ivrKnowledgeBase/findPage` | `{report.get('apiCodes', {}).get('kbFindPage')}` |
| `POST /ivrKnowledgeBase/findById/{{id}}` | `{report.get('apiCodes', {}).get('kbDetailOkCount')}` 条成功 |

## 页面抽查

页面抽查待浏览器会话补充。当前 v0.1 验收以页面中的 `试听` / `重新合成` 控件作为系统 TTS UI 证据，以节点抽屉中 `匹配知识库` 已勾选且显示具体知识库、`高级设置 -> 大模型意图分析2.0` 已勾选作为 NLP / 2.0 UI 证据。

## 失败项

{failure_text}

## 需要人工确认的字段

- 本次回归会自动开启普通节点 `matchKnowledgeBaseEnabled`，设置 `knowledgeBaseMatchType=2`，并显式绑定当前 IVR 全部知识库。
- 本次回归写入知识库答案为代表性样本，不代表完整项目交付的全量知识库导入。
- 本次回归会通过 `updateModelIntentRecognitionConfig` 启用话术级高级设置大模型意图分析 2.0，并同步保留 / 补齐普通节点级 2.0 配置；短信、转人工、黑名单、信息采集、DTMF、坐席接管仍不自动开启。
"""
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(md, encoding="utf-8")


def run_online(args: argparse.Namespace, parsed: dict[str, Any], token_text: str) -> dict[str, Any]:
    backend = resolve_backend(token_text, args.backend_region, args.backend_url)
    client = Client(backend["token"], backend)
    account = {"code": backend["accountInfo"].get("code"), "data": backend["accountInfo"].get("data") or {}}
    Client.assert_ok(account, "validate token")
    voices = client.get("/ivr/findAllTtsVoiceBaseInfo")
    Client.assert_ok(voices, "read voice list")
    voice = next((item for item in voices.get("data") or [] if int(item.get("id") or 0) == args.tts_voice_id), None) or {}

    ivr_name = args.ivr_name or f"TalkTrack-Master_v0.1_活动邀约系统TTS回归_{now_stamp()}"
    ivr_id, create_result = create_test_ivr(client, ivr_name, args.tts_voice_id, args.speech_rate)
    before_snapshot = client.get(f"/ivr/findSceneList/{ivr_id}", timeout=120)
    scene_readback, node_tts, scene_meta = configure_scene(client, ivr_id, parsed, args.tts_voice_id, args.speech_rate)
    kb_inserts, kb_page, kb_details = create_knowledge_bases(client, ivr_id, parsed, args.tts_voice_id, args.speech_rate)
    scene_readback, kb_match_meta = apply_knowledge_base_matching(client, ivr_id)
    scene_meta.update(kb_match_meta)
    model_intent_meta = update_ivr_model_intent_recognition_config(client, ivr_id)
    scene_meta.update(model_intent_meta)
    scene_readback = client.get(f"/ivr/findSceneList/{ivr_id}", timeout=120)
    Client.assert_ok(scene_readback, "read back IVR model intent recognition 2.0")
    kb_tts = [item["tts"] for item in kb_inserts]
    validation = validate_run(ivr_id, parsed, scene_readback, node_tts, kb_tts, kb_details, scene_meta)

    stamp = today_slug()
    report_path = args.report_dir / f"TalkTrack-Master_v0.1_系统TTS回归报告_{stamp}_ivr{ivr_id}.md"
    json_path = args.report_dir / f"TalkTrack-Master_v0.1_系统TTS回归_readback_{stamp}_ivr{ivr_id}.json"
    report = {
        "generatedAt": dt.datetime.now().isoformat(timespec="seconds"),
        "mode": "online",
        "backend": {
            "region": backend["region"],
            "apiBase": backend["apiBase"],
            "webBase": backend["webBase"],
        },
        "ivrId": ivr_id,
        "ivrName": ivr_name,
        "scriptGraphUrl": client.script_graph_url(ivr_id),
        "sourceDocx": parsed.get("sourcePath"),
        "voice": {
            "ttsVoiceId": args.tts_voice_id,
            "voiceName": voice.get("voiceName") or voice.get("name") or "1-普通话女声",
            "speechRate": args.speech_rate,
        },
        "parsed": parsed,
        "apiCodes": {
            "accountFindInfo": account.get("code"),
            "voiceList": voices.get("code"),
            "ivrInsert": create_result["insert"].get("code"),
            "ivrUpdate": create_result["update"].get("code"),
            "templateRead": "0",
            "updateSceneList": scene_meta["updateSceneList"].get("code"),
            "updateSceneListForKnowledgeBaseMatching": scene_meta["updateSceneListForKnowledgeBaseMatching"].get("code"),
            "updateModelIntentRecognitionConfig": scene_meta["updateModelIntentRecognitionConfig"].get("code"),
            "nodeTtsCreateOkCount": sum(1 for item in node_tts if str(item.get("createCode")) in ("0", "5")),
            "nodeTtsQueryOkCount": sum(1 for item in node_tts if str(item.get("queryCode")) == "0" and item.get("recordFilePath")),
            "kbTtsCreateOkCount": sum(1 for item in kb_tts if str(item.get("createCode")) in ("0", "5")),
            "kbTtsQueryOkCount": sum(1 for item in kb_tts if str(item.get("queryCode")) == "0" and item.get("recordFilePath")),
            "kbInsertOkCount": sum(1 for item in kb_inserts if str(item.get("insertCode")) == "0"),
            "kbFindPage": kb_page.get("code"),
            "kbDetailOkCount": sum(1 for item in kb_details if item.get("found")),
        },
        "preWriteSnapshot": {
            "ivrId": ivr_id,
            "code": before_snapshot.get("code"),
            "sceneListLength": len(safe_json_loads((before_snapshot.get("data") or {}).get("sceneList"), [])),
            "sceneListFrontendLength": len(safe_json_loads((before_snapshot.get("data") or {}).get("sceneListFrontend"), [])),
        },
        "ttsResults": {
            "nodes": node_tts,
            "knowledgeBaseAnswers": kb_tts,
        },
        "knowledgeBaseInsertResults": kb_inserts,
        "validation": validation,
        "reportPath": str(report_path),
        "jsonPath": str(json_path),
    }
    write_json(json_path, report)
    write_report(report, report_path, json_path)
    return report


def run_offline(args: argparse.Namespace, parsed: dict[str, Any], reason: str) -> dict[str, Any]:
    stamp = today_slug()
    report_path = args.report_dir / f"TalkTrack-Master_v0.1_离线解析报告_{stamp}.md"
    json_path = args.report_dir / f"TalkTrack-Master_v0.1_离线解析_readback_{stamp}.json"
    report = {
        "generatedAt": dt.datetime.now().isoformat(timespec="seconds"),
        "mode": "offline",
        "offlineReason": reason,
        "ivrId": None,
        "ivrName": None,
        "scriptGraphUrl": None,
        "sourceDocx": parsed.get("sourcePath"),
        "voice": {"ttsVoiceId": None, "voiceName": None, "speechRate": None},
        "parsed": parsed,
        "apiCodes": {},
        "validation": {"passed": False, "failureItems": [reason], "nodeSummary": [], "jumpChecks": [], "knowledgeBaseSummary": []},
        "reportPath": str(report_path),
        "jsonPath": str(json_path),
    }
    write_json(json_path, report)
    write_report(report, report_path, json_path)
    return report


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Validate TalkTrack-Master v0.1 system TTS IVR flow.")
    parser.add_argument("--docx", required=True, type=Path, help="Source DOCX path.")
    parser.add_argument("--token-env", default="SD_ADMIN_TOKEN", help="Environment variable that contains the admin token.")
    parser.add_argument("--backend-region", choices=["auto", "domestic", "overseas"], default="auto", help="Backend selector. Auto probes domestic and overseas with the provided token.")
    parser.add_argument("--backend-url", default="", help="Optional page/API URL hint such as https://ai.tbot360.com/script-graph?ivrId=3171.")
    parser.add_argument("--report-dir", default=str(DEFAULT_VAULT_REPORT_DIR), type=Path, help="Obsidian report directory.")
    parser.add_argument("--attachment-dir", default=str(DEFAULT_ATTACHMENT_DIR), type=Path, help="Obsidian attachment directory.")
    parser.add_argument("--max-kb", default=8, type=int, help="Number of representative KB answers to write in first-round regression.")
    parser.add_argument("--tts-voice-id", default=DEFAULT_TTS_VOICE_ID, type=int, help="System TTS voice id.")
    parser.add_argument("--speech-rate", default=DEFAULT_SPEECH_RATE, type=float, help="System TTS speech rate.")
    parser.add_argument("--ivr-name", default="", help="Optional test IVR name.")
    return parser.parse_args()


def main() -> int:
    sys.stdout.reconfigure(encoding="utf-8")
    args = parse_args()
    if not args.docx.exists():
        raise SystemExit(f"DOCX not found: {args.docx}")
    args.report_dir.mkdir(parents=True, exist_ok=True)
    args.attachment_dir.mkdir(parents=True, exist_ok=True)
    parsed = parse_docx(args.docx, args.max_kb)
    token = os.environ.get(args.token_env, "").strip()
    if not token:
        report = run_offline(args, parsed, f"环境变量 {args.token_env} 未提供；只完成离线解析，未写后台。")
    else:
        report = run_online(args, parsed, token)
    print(json.dumps({
        "mode": report.get("mode"),
        "passed": (report.get("validation") or {}).get("passed"),
        "ivrId": report.get("ivrId"),
        "ivrName": report.get("ivrName"),
        "reportPath": report.get("reportPath"),
        "jsonPath": report.get("jsonPath"),
        "failureItems": (report.get("validation") or {}).get("failureItems"),
        "kbCandidates": parsed.get("knowledgeBaseCandidateCount"),
        "kbWritten": parsed.get("knowledgeBaseCandidatesSelectedCount"),
    }, ensure_ascii=False, indent=2))
    return 0 if (report.get("mode") == "offline" or (report.get("validation") or {}).get("passed")) else 2


if __name__ == "__main__":
    raise SystemExit(main())
